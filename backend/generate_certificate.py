"""
TrustCode AI Compliance Certificate Generator
Generates a professional Word document certificate from audit results.
"""

import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


class CertificateGenerator:
    """Generates a professional Word certificate from audit results."""
    
    # Color scheme - Corporate Dark / Minimalist Academic
    COLORS = {
        'primary': RGBColor(0x02, 0x06, 0x17),  # Deep Slate
        'accent': RGBColor(0x22, 0xD3, 0xEE),   # Cyan
        'danger': RGBColor(0xF4, 0x3F, 0x5E),   # Rose
        'success': RGBColor(0x10, 0xB9, 0x81),  # Emerald
        'warning': RGBColor(0xF5, 0x9E, 0x0B),  # Amber
        'white': RGBColor(0xFF, 0xFF, 0xFF),
        'light_gray': RGBColor(0x94, 0xA3, 0xB8),
        'dark_gray': RGBColor(0x47, 0x55, 0x69),
    }
    
    def __init__(self, audit_json_path: str):
        self.audit_data = self._load_audit_data(audit_json_path)
        self.doc = Document()
        self._setup_document()
    
    def _load_audit_data(self, path: str) -> dict:
        """Load and validate audit JSON data."""
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if 'TrustScore' not in data:
            raise ValueError("Invalid audit data: missing 'TrustScore'")
        
        return data
    
    def _setup_document(self):
        """Configure document settings."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = self.COLORS['primary']
        
        # Set margins
        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
    
    def _add_horizontal_line(self):
        """Add a horizontal line to the document."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr {}>'
            '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="020617"/>'
            '</w:pBdr>'.format(nsdecls('w'))
        )
        pPr.append(pBdr)
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Set cell background color."""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _set_cell_text(self, cell, text: str, bold: bool = False, 
                       color: RGBColor = None, size: int = 10, alignment=None):
        """Set cell text with formatting."""
        cell.text = ''
        p = cell.paragraphs[0]
        if alignment:
            p.alignment = alignment
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        # Remove cell borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
    
    def generate(self, output_path: str = "TrustCode_Certificate.docx"):
        """Generate the complete certificate."""
        self._add_header()
        self._add_horizontal_line()
        self._add_audit_summary()
        self._add_findings_table()
        self._add_recommendation()
        self._add_footer()
        
        self.doc.save(output_path)
        print(f"[OK] Certificate saved to {output_path}")
        return output_path
    
    def _add_header(self):
        """Add the certificate header with title."""
        # Add some spacing
        self.doc.add_paragraph()
        
        # Main title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TRUSTCODE AI")
        run.font.name = 'Calibri'
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = self.COLORS['primary']
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("COMPLIANCE CERTIFICATE")
        run.font.name = 'Calibri'
        run.font.size = Pt(24)
        run.font.color.rgb = self.COLORS['accent']
        
        # Certificate ID
        cert_id = self.doc.add_paragraph()
        cert_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cert_id.add_run(f"Certificate ID: TC-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = self.COLORS['light_gray']
        
        self.doc.add_paragraph()
    
    def _add_audit_summary(self):
        """Add the audit summary table."""
        # Section title
        section_title = self.doc.add_paragraph()
        run = section_title.add_run("AUDIT SUMMARY")
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.COLORS['primary']
        
        # Create summary table
        trust_score = self.audit_data.get('TrustScore', 0)
        metadata = self.audit_data.get('AuditMetadata', {})
        
        # Determine score color
        if trust_score >= 80:
            score_color = self.COLORS['success']
            score_label = "EXCELLENT"
        elif trust_score >= 60:
            score_color = self.COLORS['accent']
            score_label = "GOOD"
        elif trust_score >= 40:
            score_color = self.COLORS['warning']
            score_label = "MODERATE"
        else:
            score_color = self.COLORS['danger']
            score_label = "CRITICAL"
        
        table = self.doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the table
        table.autofit = True
        
        # Row 0: TrustScore
        self._set_cell_shading(table.cell(0, 0), "F1F5F9")
        self._set_cell_text(table.cell(0, 0), "TrustScore", bold=True, 
                           color=self.COLORS['primary'], size=11)
        self._set_cell_shading(table.cell(0, 1), "F1F5F9")
        self._set_cell_text(table.cell(0, 1), f"{trust_score}/100 - {score_label}", 
                           bold=True, color=score_color, size=14)
        
        # Row 1: Audit Date
        self._set_cell_text(table.cell(1, 0), "Audit Date", bold=True, 
                           color=self.COLORS['primary'], size=11)
        audit_date = metadata.get('audit_date', 'N/A')
        self._set_cell_text(table.cell(1, 1), audit_date, color=self.COLORS['dark_gray'], size=11)
        
        # Row 2: File Analyzed
        self._set_cell_text(table.cell(2, 0), "File Analyzed", bold=True, 
                           color=self.COLORS['primary'], size=11)
        file_name = metadata.get('file', 'N/A')
        self._set_cell_text(table.cell(2, 1), file_name, color=self.COLORS['dark_gray'], size=11)
        
        # Row 3: Total Findings
        self._set_cell_text(table.cell(3, 0), "Total Findings", bold=True, 
                           color=self.COLORS['primary'], size=11)
        total = metadata.get('total_findings', 0)
        self._set_cell_text(table.cell(3, 1), str(total), color=self.COLORS['dark_gray'], size=11)
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(12)
        
        self.doc.add_paragraph()
    
    def _add_findings_table(self):
        """Add the detailed findings table."""
        findings = self.audit_data.get('Findings', [])
        
        if not findings:
            return
        
        # Section title
        section_title = self.doc.add_paragraph()
        run = section_title.add_run("DETAILED FINDINGS")
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.COLORS['primary']
        
        # Create findings table
        table = self.doc.add_table(rows=1 + len(findings), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Severity', 'Category', 'Issue', 'Line']
        for i, header in enumerate(headers):
            self._set_cell_shading(table.cell(0, i), "020617")
            self._set_cell_text(table.cell(0, i), header, bold=True, 
                               color=self.COLORS['white'], size=9)
        
        # Data rows
        severity_colors = {
            'critical': self.COLORS['danger'],
            'high': self.COLORS['danger'],
            'medium': self.COLORS['warning'],
            'low': self.COLORS['accent'],
        }
        
        for idx, finding in enumerate(findings):
            row_idx = idx + 1
            
            # Severity
            severity = finding.get('severity', 'low').upper()
            self._set_cell_text(table.cell(row_idx, 0), severity, bold=True,
                               color=severity_colors.get(finding.get('severity', 'low'), 
                                                        self.COLORS['dark_gray']),
                               size=8)
            
            # Category
            self._set_cell_text(table.cell(row_idx, 1), finding.get('category', 'N/A'),
                               color=self.COLORS['dark_gray'], size=8)
            
            # Issue (truncated)
            message = finding.get('message', '')
            if len(message) > 60:
                message = message[:57] + '...'
            self._set_cell_text(table.cell(row_idx, 2), message,
                               color=self.COLORS['primary'], size=8)
            
            # Line
            self._set_cell_text(table.cell(row_idx, 3), str(finding.get('line', 'N/A')),
                               color=self.COLORS['dark_gray'], size=8,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.doc.add_paragraph()
    
    def _add_recommendation(self):
        """Add the PhD-level recommendation."""
        recommendation = self.audit_data.get('PhD_Level_Recommendation', '')
        
        if not recommendation:
            return
        
        # Section title
        section_title = self.doc.add_paragraph()
        run = section_title.add_run("PHD-LEVEL RECOMMENDATION")
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.COLORS['primary']
        
        # Recommendation text
        rec_para = self.doc.add_paragraph()
        run = rec_para.add_run(recommendation)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = self.COLORS['dark_gray']
        run.font.italic = True
        
        self.doc.add_paragraph()
    
    def _add_footer(self):
        """Add the certification footer."""
        self._add_horizontal_line()
        
        # Footer text
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("CERTIFIED BY TRUSTCODE AI ENGINE")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = self.COLORS['primary']
        
        # PhD badge
        badge = self.doc.add_paragraph()
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = badge.add_run("PHD RESEARCH STANDARDS")
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = self.COLORS['accent']
        
        # Creator credit
        credit = self.doc.add_paragraph()
        credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = credit.add_run("Created by Ali Hasan")
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = self.COLORS['dark_gray']
        
        portfolio = self.doc.add_paragraph()
        portfolio.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = portfolio.add_run("https://alool266.github.io/portfolio-website/")
        run.font.name = 'Calibri'
        run.font.size = Pt(7)
        run.font.color.rgb = self.COLORS['accent']
        
        # Disclaimer
        disclaimer = self.doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run(
            "This certificate is generated automatically by TrustCode AI Audit Engine v1.0.0. "
            "It represents a static analysis assessment and should be used as a guideline, "
            "not a guarantee of code quality."
        )
        run.font.name = 'Calibri'
        run.font.size = Pt(7)
        run.font.color.rgb = self.COLORS['light_gray']


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print("Usage: python generate_certificate.py <audit_results.json> [output.docx]")
        sys.exit(1)
    
    audit_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "TrustCode_Certificate.docx"
    
    generator = CertificateGenerator(audit_path)
    generator.generate(output_path)


if __name__ == "__main__":
    main()